"""
Script de génération de l'étude technique détaillée du projet E-Carto
Génère un fichier Word avec tous les détails techniques incluant la documentation BDD complète
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime
import os

def set_cell_shading(cell, fill_color):
    """Appliquer une couleur de fond à une cellule"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), fill_color)
    cell._tc.get_or_add_tcPr().append(shading_elm)

def create_styled_table(doc, headers, rows, header_color='2E74B5'):
    """Créer un tableau stylisé"""
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    
    # En-têtes
    header_row = table.rows[0]
    for i, header in enumerate(headers):
        cell = header_row.cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        set_cell_shading(cell, header_color)
    
    # Données
    for row_data in rows:
        row = table.add_row()
        for i, value in enumerate(row_data):
            row.cells[i].text = str(value)
    
    return table

def add_uml_diagram_text(doc):
    """Ajouter le diagramme UML en notation textuelle PlantUML"""
    uml_code = """
┌─────────────────────────────────────────────────────────────────────────────────────────┐
│                           DIAGRAMME DE CLASSES UML - E-CARTO                            │
└─────────────────────────────────────────────────────────────────────────────────────────┘

        ┌──────────────────────┐
        │   <<abstract>>       │
        │     AuditModel       │
        ├──────────────────────┤
        │ - created_at         │
        │ - updated_at         │
        │ - created_by         │
        │ - updated_by         │
        └──────────────────────┘
                   △
                   │ extends
     ┌─────────────┼─────────────┬─────────────┐
     │             │             │             │
     ▼             ▼             ▼             ▼
┌─────────┐  ┌──────────┐  ┌─────────┐  ┌───────────┐
│ Projet  │  │   Site   │  │ Rapport │  │TypeProjet │
└─────────┘  └──────────┘  └─────────┘  └───────────┘


                        RELATIONS PRINCIPALES
                        ═════════════════════

┌──────────────────┐         1..* ┌──────────────────┐
│    Utilisateur   │◆─────────────│   Notification   │
└──────────────────┘              └──────────────────┘
        │
        │ 1
        ▼
┌──────────────────┐
│   utilisateur_   │
│   permissions    │
└──────────────────┘


┌──────────────────┐   *..1  ┌──────────────────┐
│      Projet      │────────▶│    TypeProjet    │
│                  │         └──────────────────┘
│                  │
│                  │   *..1  ┌──────────────────┐
│                  │────────▶│       Site       │
└──────────────────┘         └──────────────────┘
        │
        │ 1..*
        ▼
┌──────────────────┐
│   projet_equipe  │
│    (membres)     │
└──────────────────┘


┌──────────────────┐   *..1  ┌──────────────────┐
│     Rapport      │────────▶│      Projet      │
└──────────────────┘         └──────────────────┘
"""
    
    para = doc.add_paragraph()
    run = para.add_run(uml_code)
    run.font.name = 'Consolas'
    run.font.size = Pt(8)

def generate_etude_technique():
    doc = Document()
    
    # Configuration des styles
    style = doc.styles['Heading 1']
    style.font.size = Pt(18)
    style.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
    
    style = doc.styles['Heading 2']
    style.font.size = Pt(14)
    style.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
    
    style = doc.styles['Heading 3']
    style.font.size = Pt(12)
    style.font.color.rgb = RGBColor(0x40, 0x40, 0x40)
    
    # === PAGE DE TITRE ===
    title = doc.add_paragraph()
    title_run = title.add_run("ÉTUDE TECHNIQUE DÉTAILLÉE")
    title_run.font.size = Pt(28)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    subtitle = doc.add_paragraph()
    subtitle_run = subtitle.add_run("Projet E-Carto")
    subtitle_run.font.size = Pt(24)
    subtitle_run.font.italic = True
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    desc = doc.add_paragraph()
    desc_run = desc.add_run("Plateforme de Gestion des Projets Informatiques\nEranove Academy")
    desc_run.font.size = Pt(16)
    desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    for _ in range(3):
        doc.add_paragraph()
    
    date_para = doc.add_paragraph()
    date_run = date_para.add_run(f"Date: {datetime.date.today().strftime('%d/%m/%Y')}")
    date_run.font.size = Pt(12)
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_page_break()
    
    # === TABLE DES MATIÈRES ===
    doc.add_heading("TABLE DES MATIÈRES", 1)
    toc_items = [
        "1. Présentation du Projet",
        "2. Architecture Technique",
        "   2.1. Architecture Globale",
        "   2.2. Stack Technologique",
        "3. Backend - Spring Boot",
        "   3.1. Structure du Projet",
        "   3.2. Modules Fonctionnels",
        "   3.3. Sécurité et Authentification",
        "   3.4. Configuration",
        "4. Frontend - Angular",
        "   4.1. Architecture Frontend",
        "   4.2. Modules et Composants",
        "   4.3. Cartographie Leaflet",
        "5. API REST",
        "   5.1. Endpoints Principaux",
        "   5.2. Documentation OpenAPI",
        "6. Documentation Base de Données",
        "   6.1. Modèle Conceptuel de Données (MCD)",
        "   6.2. Diagramme UML des Entités",
        "   6.3. Description des Tables",
        "   6.4. Relations entre Entités",
        "   6.5. Dictionnaire de Données",
        "7. Sécurité",
        "8. Recommandations",
    ]
    for item in toc_items:
        doc.add_paragraph(item)
    
    doc.add_page_break()
    
    # === 1. PRÉSENTATION DU PROJET ===
    doc.add_heading("1. Présentation du Projet", 1)
    
    doc.add_heading("1.1. Contexte", 2)
    doc.add_paragraph(
        "Le projet E-Carto est une plateforme de gestion des projets informatiques développée "
        "dans le cadre du stage Eranove Academy. Cette application web full-stack permet de gérer "
        "l'ensemble du cycle de vie des projets informatiques, incluant la cartographie des sites, "
        "la gestion des rapports de faisabilité et le suivi des équipes."
    )
    
    doc.add_heading("1.2. Objectifs", 2)
    objectives = [
        "Centraliser la gestion des projets informatiques",
        "Visualiser géographiquement les sites et projets sur une carte interactive",
        "Gérer les rapports de faisabilité et d'analyse de risques",
        "Assurer une gestion fine des utilisateurs et de leurs rôles",
        "Exporter des rapports au format PDF",
        "Permettre la recherche avancée et le filtrage des données"
    ]
    for obj in objectives:
        doc.add_paragraph(f"• {obj}")
    
    doc.add_heading("1.3. Périmètre Fonctionnel", 2)
    modules = [
        ("Gestion des Projets", "CRUD complet, suivi de progression, gestion des équipes"),
        ("Gestion des Sites", "Géolocalisation, cartographie, informations de contact"),
        ("Rapports de Faisabilité", "Analyse de risques, estimation budgétaire, recommandations"),
        ("Types de Projet", "Catégorisation et classification des projets"),
        ("Gestion des Utilisateurs", "Authentification, autorisation, gestion des rôles"),
        ("Dashboard", "Tableaux de bord, statistiques, KPIs")
    ]
    
    table = create_styled_table(doc, ["Module", "Description"], modules)
    
    doc.add_page_break()
    
    # === 2. ARCHITECTURE TECHNIQUE ===
    doc.add_heading("2. Architecture Technique", 1)
    
    doc.add_heading("2.1. Architecture Globale", 2)
    doc.add_paragraph(
        "L'application suit une architecture client-serveur moderne avec séparation claire "
        "entre le frontend et le backend. Le frontend Angular communique avec le backend Spring Boot "
        "via une API REST sécurisée par JWT."
    )
    
    arch_layers = [
        ("Présentation", "Angular 17, Material Design, Leaflet.js"),
        ("API Gateway", "Spring Boot REST Controllers, OpenAPI/Swagger"),
        ("Logique Métier", "Services Spring, MapStruct, Validation"),
        ("Persistance", "Spring Data JPA, Hibernate"),
        ("Base de Données", "PostgreSQL avec Flyway migrations")
    ]
    create_styled_table(doc, ["Couche", "Technologies"], arch_layers)
    
    doc.add_paragraph()
    doc.add_heading("2.2. Stack Technologique", 2)
    
    doc.add_heading("Backend", 3)
    backend_tech = [
        ("Framework", "Spring Boot 3.2.0"),
        ("Langage", "Java 17 (LTS)"),
        ("ORM", "Hibernate / Spring Data JPA"),
        ("Sécurité", "Spring Security 6 + JWT (jjwt 0.12.3)"),
        ("Base de données", "PostgreSQL"),
        ("Migrations", "Flyway"),
        ("Mapping DTO", "MapStruct 1.5.5"),
        ("Génération de code", "Lombok 1.18.30"),
        ("Documentation API", "SpringDoc OpenAPI 2.2.0"),
        ("Génération PDF", "OpenPDF 1.3.30"),
        ("Analyse de fichiers", "Apache Tika 2.9.0"),
        ("Envoi d'emails", "Spring Mail"),
        ("Tests", "JUnit 5, TestContainers")
    ]
    create_styled_table(doc, ["Composant", "Version/Technologie"], backend_tech)
    
    doc.add_paragraph()
    doc.add_heading("Frontend", 3)
    frontend_tech = [
        ("Framework", "Angular 17.3.0"),
        ("Langage", "TypeScript 5.2"),
        ("UI Components", "Angular Material 17.3.0"),
        ("CDK", "Angular CDK 17.3.0"),
        ("Cartographie", "Leaflet 1.9.4"),
        ("CSS", "TailwindCSS 3.3.0"),
        ("Gestion d'état", "RxJS 7.8"),
        ("Build Tool", "Angular CLI 17.3.0")
    ]
    create_styled_table(doc, ["Composant", "Version"], frontend_tech)
    
    doc.add_page_break()
    
    # === 3. BACKEND - SPRING BOOT ===
    doc.add_heading("3. Backend - Spring Boot", 1)
    
    doc.add_heading("3.1. Structure du Projet", 2)
    doc.add_paragraph(
        "Le backend suit une architecture modulaire organisée par domaine fonctionnel. "
        "Chaque module contient ses propres controllers, services, repositories, DTOs et mappers."
    )
    
    structure = """
com.gs2e.stage_eranove_academy/
├── StageEranoveAcademyApplication.java  (Point d'entrée)
├── common/                              (Utilitaires partagés)
│   ├── model/                           (AuditModel, entités de base)
│   └── Exceptions/                      (Exceptions personnalisées)
├── config/                              (Configuration Spring)
│   ├── SecurityConfig.java
│   └── JacksonConfig.java
├── projet/                              (Module Projets)
│   ├── controller/
│   ├── dto/
│   ├── mapper/
│   ├── model/
│   ├── repository/
│   ├── service/
│   └── validator/
├── rapport/                             (Module Rapports)
├── security/                            (Module Sécurité)
│   ├── controller/
│   ├── dto/
│   ├── model/
│   ├── repository/
│   ├── service/
│   └── util/
├── site/                                (Module Sites)
└── typeprojet/                          (Module Types de Projet)
"""
    doc.add_paragraph(structure, style='No Spacing')
    
    doc.add_heading("3.2. Modules Fonctionnels", 2)
    
    doc.add_heading("3.2.1. Module Projet", 3)
    doc.add_paragraph(
        "Le module Projet gère l'ensemble du cycle de vie des projets avec un contrôleur "
        "complet exposant plus de 20 endpoints REST."
    )
    
    projet_endpoints = [
        ("GET", "/api/projets", "Liste paginée des projets"),
        ("GET", "/api/projets/{id}", "Détails d'un projet"),
        ("POST", "/api/projets", "Création d'un projet"),
        ("PUT", "/api/projets/{id}", "Mise à jour d'un projet"),
        ("DELETE", "/api/projets/{id}", "Suppression d'un projet"),
        ("GET", "/api/projets/search", "Recherche avancée multicritères"),
        ("GET", "/api/projets/status/{status}", "Projets par statut"),
        ("GET", "/api/projets/responsable/{resp}", "Projets par responsable"),
        ("GET", "/api/projets/type/{typeId}", "Projets par type"),
        ("GET", "/api/projets/retard", "Projets en retard"),
        ("GET", "/api/projets/statistiques", "Statistiques globales"),
        ("PATCH", "/api/projets/{id}/status", "Mise à jour du statut"),
        ("PATCH", "/api/projets/{id}/progress", "Mise à jour de la progression"),
        ("GET", "/api/projets/{id}/report", "Export PDF du projet"),
        ("POST", "/api/projets/{id}/send-report", "Envoi du rapport par email")
    ]
    create_styled_table(doc, ["Méthode", "Endpoint", "Description"], projet_endpoints)
    
    doc.add_page_break()
    
    doc.add_heading("3.3. Sécurité et Authentification", 2)
    
    doc.add_paragraph(
        "L'application utilise une authentification basée sur JWT (JSON Web Tokens) "
        "avec Spring Security 6. Les sessions sont stateless pour une scalabilité optimale."
    )
    
    doc.add_heading("3.3.1. Configuration Spring Security", 3)
    security_config = [
        ("Authentification", "JWT avec Bearer Token"),
        ("Encodage mot de passe", "BCrypt (strength: 8)"),
        ("Gestion de session", "Stateless"),
        ("CORS", "Configuré pour localhost:*"),
        ("CSRF", "Désactivé (API REST)"),
        ("Durée token JWT", "24 heures (86400000 ms)"),
        ("Refresh Token", "7 jours (604800000 ms)")
    ]
    create_styled_table(doc, ["Paramètre", "Configuration"], security_config)
    
    doc.add_heading("3.3.2. Système de Rôles", 3)
    doc.add_paragraph("L'application implémente 6 rôles utilisateur avec des permissions différenciées:")
    
    roles = [
        ("ADMINISTRATEUR_SYSTEME", "Accès complet à toutes les fonctionnalités", "Tous les modules"),
        ("CHEF_DE_PROJET", "Gestion des projets et types de projets", "Projets, Sites, Rapports, Types"),
        ("ANALYSTE", "Consultation et analyse des projets", "Projets, Sites, Rapports"),
        ("DEVELOPPEUR", "Consultation des projets assignés", "Sites, Rapports"),
        ("DECIDEUR", "Vue direction, statistiques", "Sites, Rapports"),
        ("OBSERVATEUR", "Lecture seule", "Sites, Rapports")
    ]
    create_styled_table(doc, ["Rôle", "Description", "Accès Modules"], roles)
    
    doc.add_page_break()
    
    # === 4. FRONTEND - ANGULAR ===
    doc.add_heading("4. Frontend - Angular", 1)
    
    doc.add_heading("4.1. Architecture Frontend", 2)
    doc.add_paragraph(
        "Le frontend est construit avec Angular 17 en mode standalone, "
        "utilisant les dernières fonctionnalités comme les signaux et le lazy loading avancé."
    )
    
    frontend_structure = """
src/app/
├── app.component.ts           (Composant racine)
├── app.config.ts              (Configuration standalone)
├── app.routes.ts              (Routes principales)
├── auth/                      (Module authentification)
│   ├── login/                 (Page de connexion)
│   ├── guards/                (AuthGuard, RoleGuard)
│   └── services/              (AuthService)
├── features/                  (Modules fonctionnels)
│   ├── dashboard-home/        (Accueil dashboard)
│   ├── projets/               (Gestion projets)
│   ├── rapports/              (Gestion rapports)
│   ├── sites/                 (Gestion sites)
│   ├── types-projet/          (Types de projet)
│   └── utilisateurs/          (Gestion utilisateurs)
├── layout/                    (Composants de mise en page)
│   └── dashboard/             (Layout principal)
└── shared/                    (Éléments partagés)
    ├── components/            (Composants réutilisables)
    ├── services/              (Services partagés)
    ├── models/                (Interfaces TypeScript)
    └── enums/                 (Énumérations)
"""
    doc.add_paragraph(frontend_structure, style='No Spacing')
    
    doc.add_heading("4.2. Modules et Composants", 2)
    
    features = [
        ("Dashboard Home", "Tableau de bord principal avec statistiques et KPIs"),
        ("Projets", "Liste, création, édition, détails des projets"),
        ("Rapports", "Upload et gestion des rapports de faisabilité"),
        ("Sites", "Gestion des sites avec carte interactive"),
        ("Types Projet", "Administration des catégories de projets"),
        ("Utilisateurs", "Gestion des comptes et rôles")
    ]
    create_styled_table(doc, ["Feature Module", "Description"], features)
    
    doc.add_heading("4.3. Cartographie Leaflet", 2)
    doc.add_paragraph(
        "L'application intègre Leaflet 1.9.4 pour l'affichage d'une carte interactive "
        "centrée sur la Côte d'Ivoire. Les sites sont représentés par des marqueurs avec "
        "popups informatifs."
    )
    
    doc.add_page_break()
    
    # === 5. API REST ===
    doc.add_heading("5. API REST", 1)
    
    doc.add_heading("5.1. Conventions", 2)
    conventions = [
        ("Base URL", "/api"),
        ("Format", "JSON"),
        ("Authentification", "Bearer Token (Header Authorization)"),
        ("Pagination", "Page-based (page, size, sort)"),
        ("Codes HTTP", "200, 201, 204, 400, 401, 403, 404, 500")
    ]
    create_styled_table(doc, ["Aspect", "Convention"], conventions)
    
    doc.add_heading("5.2. Documentation OpenAPI", 2)
    doc.add_paragraph(
        "L'API est documentée via SpringDoc OpenAPI et accessible via Swagger UI à l'URL: "
        "/swagger-ui.html. La spécification complète est disponible à /v3/api-docs."
    )
    
    doc.add_page_break()
    
    # === 6. DOCUMENTATION BASE DE DONNÉES ===
    doc.add_heading("6. Documentation Base de Données", 1)
    
    doc.add_paragraph(
        "Cette section présente la documentation complète de la base de données PostgreSQL "
        "utilisée par l'application E-Carto, incluant le modèle conceptuel, les diagrammes UML, "
        "les relations entre entités et le dictionnaire de données."
    )
    
    # 6.1 MCD
    doc.add_heading("6.1. Modèle Conceptuel de Données (MCD)", 2)
    doc.add_paragraph(
        "Le modèle de données suit une approche relationnelle normalisée. Les entités principales "
        "sont organisées autour du concept central de 'Projet', qui est lié aux Sites, Types de Projet, "
        "Rapports et Utilisateurs."
    )
    
    # Schéma conceptuel
    mcd_entities = [
        ("UTILISATEUR", "Gère les comptes, authentification et autorisations"),
        ("PROJET", "Entité centrale représentant un projet informatique"),
        ("SITE", "Localisation géographique avec coordonnées GPS"),
        ("TYPE_PROJET", "Catégorie/classification des projets"),
        ("RAPPORT", "Rapports de faisabilité attachés aux projets"),
        ("NOTIFICATION", "Alertes et messages pour les utilisateurs"),
        ("PERMISSION", "Matrice de droits d'accès par ressource")
    ]
    create_styled_table(doc, ["Entité", "Description"], mcd_entities)
    
    doc.add_paragraph()
    
    # 6.2 Diagramme UML
    doc.add_heading("6.2. Diagramme UML des Entités", 2)
    doc.add_paragraph(
        "Le diagramme suivant présente les classes d'entités JPA et leurs relations. "
        "Toutes les entités principales héritent de la classe abstraite AuditModel "
        "qui fournit les champs d'audit (created_at, updated_at, created_by, updated_by)."
    )
    
    add_uml_diagram_text(doc)
    
    doc.add_page_break()
    
    # 6.3 Description détaillée des tables
    doc.add_heading("6.3. Description Détaillée des Tables", 2)
    
    # Table UTILISATEURS
    doc.add_heading("6.3.1. Table: utilisateurs", 3)
    doc.add_paragraph("Stocke les informations des comptes utilisateurs et leurs credentials.")
    
    utilisateurs_fields = [
        ("id", "BIGSERIAL", "PK", "Identifiant unique auto-incrémenté"),
        ("email", "VARCHAR(254)", "UNIQUE, NOT NULL", "Adresse email (login)"),
        ("password", "VARCHAR(254)", "NOT NULL", "Mot de passe hashé BCrypt"),
        ("nom", "VARCHAR(100)", "NOT NULL", "Nom de famille"),
        ("prenom", "VARCHAR(100)", "NOT NULL", "Prénom"),
        ("telephone", "VARCHAR(20)", "", "Numéro de téléphone"),
        ("departement", "VARCHAR(100)", "", "Département/Service"),
        ("poste", "VARCHAR(100)", "", "Poste occupé"),
        ("role", "VARCHAR(50)", "NOT NULL", "Rôle (ENUM)"),
        ("actif", "BOOLEAN", "DEFAULT true", "Compte actif/désactivé"),
        ("created_at", "TIMESTAMP", "NOT NULL", "Date de création"),
        ("updated_at", "TIMESTAMP", "NOT NULL", "Date de modification"),
        ("refresh_token", "VARCHAR(500)", "", "Token de rafraîchissement JWT"),
        ("refresh_token_expiry", "TIMESTAMP", "", "Expiration du refresh token")
    ]
    create_styled_table(doc, ["Colonne", "Type", "Contraintes", "Description"], utilisateurs_fields)
    
    doc.add_paragraph()
    
    # Table PROJETS
    doc.add_heading("6.3.2. Table: projets", 3)
    doc.add_paragraph("Table principale contenant les informations des projets informatiques.")
    
    projets_fields = [
        ("id", "BIGSERIAL", "PK", "Identifiant unique"),
        ("nom", "VARCHAR(255)", "NOT NULL", "Nom du projet"),
        ("description", "TEXT", "", "Description détaillée"),
        ("statut", "VARCHAR(20)", "NOT NULL", "PREVU, EN_COURS, TERMINE, ANNULE"),
        ("priorite", "VARCHAR(20)", "NOT NULL", "FAIBLE, MOYENNE, HAUTE, CRITIQUE"),
        ("responsable", "VARCHAR(255)", "NOT NULL", "Nom du responsable"),
        ("date_debut", "DATE", "NOT NULL", "Date de démarrage"),
        ("date_fin_prevue", "DATE", "", "Date de fin planifiée"),
        ("date_fin_reelle", "DATE", "", "Date de fin effective"),
        ("budget", "DECIMAL(12,2)", "", "Budget alloué en FCFA"),
        ("progression", "INTEGER", "DEFAULT 0", "Pourcentage d'avancement (0-100)"),
        ("type_projet_id", "BIGINT", "FK → type_projet", "Référence au type"),
        ("site_id", "BIGINT", "FK → sites", "Référence au site"),
        ("tags", "VARCHAR(500)", "", "Étiquettes séparées par virgules"),
        ("created_at", "TIMESTAMP", "NOT NULL", "Date de création"),
        ("updated_at", "TIMESTAMP", "NOT NULL", "Date de modification"),
        ("created_by", "VARCHAR(100)", "", "Créateur"),
        ("updated_by", "VARCHAR(100)", "", "Dernier modificateur")
    ]
    create_styled_table(doc, ["Colonne", "Type", "Contraintes", "Description"], projets_fields)
    
    doc.add_page_break()
    
    # Table SITES
    doc.add_heading("6.3.3. Table: sites", 3)
    doc.add_paragraph("Contient les sites géographiques avec leurs coordonnées pour la cartographie.")
    
    sites_fields = [
        ("id", "BIGSERIAL", "PK", "Identifiant unique"),
        ("nom", "VARCHAR(255)", "NOT NULL", "Nom du site"),
        ("description", "TEXT", "", "Description"),
        ("adresse", "VARCHAR(500)", "NOT NULL", "Adresse postale"),
        ("ville", "VARCHAR(100)", "NOT NULL", "Ville"),
        ("region", "VARCHAR(100)", "NOT NULL", "Région"),
        ("pays", "VARCHAR(100)", "DEFAULT 'Côte d'Ivoire'", "Pays"),
        ("latitude", "DOUBLE", "NOT NULL", "Coordonnée GPS latitude"),
        ("longitude", "DOUBLE", "NOT NULL", "Coordonnée GPS longitude"),
        ("type", "VARCHAR(50)", "NOT NULL", "Type de site (ENUM)"),
        ("statut", "VARCHAR(50)", "DEFAULT 'ACTIF'", "Statut du site"),
        ("contact_personne", "VARCHAR(200)", "", "Personne de contact"),
        ("contact_telephone", "VARCHAR(20)", "", "Téléphone contact"),
        ("contact_email", "VARCHAR(254)", "", "Email contact"),
        ("nombre_employes", "INTEGER", "", "Effectif"),
        ("horaires_ouverture", "VARCHAR(200)", "", "Horaires"),
        ("equipements", "TEXT", "", "Liste des équipements")
    ]
    create_styled_table(doc, ["Colonne", "Type", "Contraintes", "Description"], sites_fields)
    
    doc.add_paragraph()
    
    # Table RAPPORTS
    doc.add_heading("6.3.4. Table: rapports", 3)
    doc.add_paragraph("Stocke les rapports de faisabilité associés aux projets avec analyse de risques.")
    
    rapports_fields = [
        ("id", "BIGSERIAL", "PK", "Identifiant unique"),
        ("nom", "VARCHAR(200)", "NOT NULL", "Nom du rapport"),
        ("description", "TEXT", "", "Description"),
        ("fichier_nom", "VARCHAR(255)", "NOT NULL", "Nom du fichier uploadé"),
        ("fichier_type", "VARCHAR(20)", "NOT NULL", "Type MIME"),
        ("fichier_taille", "BIGINT", "NOT NULL", "Taille en octets"),
        ("fichier_chemin", "VARCHAR(500)", "NOT NULL", "Chemin de stockage"),
        ("projet_id", "BIGINT", "FK → projets", "Projet associé"),
        ("uploade_par", "VARCHAR(100)", "NOT NULL", "Auteur de l'upload"),
        ("faisabilite", "INTEGER", "NOT NULL", "Score faisabilité (0-100)"),
        ("risque", "VARCHAR(20)", "NOT NULL", "FAIBLE, MOYEN, ELEVE, CRITIQUE"),
        ("budget_estime", "DECIMAL(12,2)", "", "Budget estimé"),
        ("duree_estimee_mois", "INTEGER", "NOT NULL", "Durée estimée en mois"),
        ("recommandations", "TEXT", "", "Recommandations"),
        ("analyse_automatique", "BOOLEAN", "DEFAULT false", "Généré automatiquement"),
        ("version", "BIGINT", "", "Version pour optimistic locking")
    ]
    create_styled_table(doc, ["Colonne", "Type", "Contraintes", "Description"], rapports_fields)
    
    doc.add_paragraph()
    doc.add_paragraph("Index définis sur la table rapports:")
    rapports_indexes = [
        ("idx_rapport_projet_id", "projet_id", "Recherche par projet"),
        ("idx_rapport_upload_date", "created_at", "Tri par date d'upload"),
        ("idx_rapport_risque", "risque", "Filtrage par niveau de risque"),
        ("idx_rapport_faisabilite", "faisabilite", "Filtrage par score")
    ]
    create_styled_table(doc, ["Nom Index", "Colonne(s)", "Utilisation"], rapports_indexes)
    
    doc.add_page_break()
    
    # Table TYPE_PROJET
    doc.add_heading("6.3.5. Table: type_projet", 3)
    doc.add_paragraph("Définit les catégories de projets avec soft delete support.")
    
    type_projet_fields = [
        ("id_type_projet", "BIGSERIAL", "PK", "Identifiant unique"),
        ("nom", "VARCHAR(100)", "", "Nom technique"),
        ("libelle", "VARCHAR(200)", "", "Libellé d'affichage"),
        ("description", "TEXT", "", "Description"),
        ("couleur", "VARCHAR(7)", "", "Code couleur hex (#RRGGBB)"),
        ("icone", "VARCHAR(50)", "", "Nom de l'icône Material"),
        ("est_actif", "BOOLEAN", "DEFAULT true", "Soft delete flag"),
        ("created_at", "TIMESTAMP", "NOT NULL", "Date de création"),
        ("updated_at", "TIMESTAMP", "NOT NULL", "Date de modification")
    ]
    create_styled_table(doc, ["Colonne", "Type", "Contraintes", "Description"], type_projet_fields)
    
    doc.add_paragraph()
    
    # Table NOTIFICATIONS
    doc.add_heading("6.3.6. Table: notifications", 3)
    doc.add_paragraph("Gère les notifications en temps réel pour les utilisateurs.")
    
    notifications_fields = [
        ("id", "BIGSERIAL", "PK", "Identifiant unique"),
        ("utilisateur_id", "BIGINT", "FK → utilisateurs, NOT NULL", "Destinataire"),
        ("type", "VARCHAR(50)", "NOT NULL", "Type de notification"),
        ("titre", "VARCHAR(200)", "NOT NULL", "Titre court"),
        ("message", "TEXT", "NOT NULL", "Contenu du message"),
        ("lu", "BOOLEAN", "DEFAULT false", "Statut de lecture"),
        ("created_at", "TIMESTAMP", "NOT NULL", "Date de création")
    ]
    create_styled_table(doc, ["Colonne", "Type", "Contraintes", "Description"], notifications_fields)
    
    doc.add_paragraph()
    
    # Table PROJET_EQUIPE
    doc.add_heading("6.3.7. Table: projet_equipe", 3)
    doc.add_paragraph("Table de jointure pour les membres d'équipe d'un projet (ElementCollection).")
    
    projet_equipe_fields = [
        ("projet_id", "BIGINT", "FK → projets, NOT NULL", "Référence au projet"),
        ("membre", "VARCHAR(255)", "NOT NULL", "Nom du membre de l'équipe")
    ]
    create_styled_table(doc, ["Colonne", "Type", "Contraintes", "Description"], projet_equipe_fields)

    doc.add_paragraph()

    doc.add_heading("6.3.8. Modèles de Workflow (Phases, Étapes)", 3)
    doc.add_paragraph("Divise le projet en phases séquentielles gérant chacune plusieurs étapes avec workflow d'approbation.")
    workflow_fields = [
        ("phases.id", "BIGSERIAL", "PK", "Identifiant de la phase"),
        ("phases.nom", "VARCHAR(150)", "NOT NULL", "Nom de la phase"),
        ("phases.progression", "INTEGER", "", "Progression de la phase en %"),
        ("etapes.id", "BIGSERIAL", "PK", "Identifiant de l'étape"),
        ("etapes.phase_id", "BIGINT", "FK → phases", "Phase parente"),
        ("etapes.statut", "VARCHAR", "NOT NULL", "A_FAIRE, EN_COURS, VALIDEE..."),
        ("etapes.bloquante", "BOOLEAN", "NOT NULL", "Si true, bloque la suite"),
        ("validations_etapes.id", "BIGSERIAL", "PK", "Trace la validation"),
        ("validations_etapes.decision", "VARCHAR", "NOT NULL", "APPROUVEE, REJETEE")
    ]
    create_styled_table(doc, ["Colonne", "Type", "Contraintes", "Description"], workflow_fields)

    doc.add_paragraph()

    doc.add_heading("6.3.9. Gouvernance (Risques, Alertes, Historique)", 3)
    doc.add_paragraph("Gestion proactive des alertes et traçabilité inviolable.")
    gov_fields = [
        ("risques.probabilite", "VARCHAR", "ENUM", "FAIBLE, MOYEN, ELEVE, CRITIQUE"),
        ("risques.impact", "VARCHAR", "ENUM", "FAIBLE, MOYEN, ELEVE, CRITIQUE"),
        ("risques.statut", "VARCHAR", "ENUM", "IDENTIFIE, TRAITE, CLOS"),
        ("alertes.niveau", "VARCHAR", "ENUM", "INFO, AVERTISSEMENT, URGENT"),
        ("alertes.lue", "BOOLEAN", "DEFAULT false", "Indice de lecture"),
        ("historique_statuts.motif", "TEXT", "", "Motif du changement"),
        ("historique_statuts.utilisateur_id", "BIGINT", "FK", "Auteur")
    ]
    create_styled_table(doc, ["Colonne", "Type", "Contraintes", "Description"], gov_fields)
    
    doc.add_page_break()
    
    # 6.4 Relations entre entités
    doc.add_heading("6.4. Relations entre Entités", 2)
    
    doc.add_paragraph(
        "Cette section décrit les relations entre les différentes entités du système. "
        "Le modèle utilise JPA/Hibernate avec les annotations appropriées pour définir ces relations."
    )
    
    doc.add_heading("6.4.1. Diagramme des Relations", 3)
    
    relations_diagram = """
┌─────────────────────────────────────────────────────────────────────────────────┐
│                        DIAGRAMME ENTITÉ-RELATION                               │
└─────────────────────────────────────────────────────────────────────────────────┘

    ┌─────────────┐                    ┌─────────────┐
    │ UTILISATEUR │ 1 ──────────── * │NOTIFICATION │
    └─────────────┘                    └─────────────┘
          │
          │ 1
          ▼
    ┌─────────────┐
    │ PERMISSIONS │
    └─────────────┘


    ┌─────────────┐                    ┌─────────────┐
    │   PROJET    │ * ──────────── 1  │ TYPE_PROJET │
    │             │                    └─────────────┘
    │             │
    │             │ * ──────────── 1  ┌─────────────┐
    │             │                    │    SITE     │
    └─────────────┘                    └─────────────┘
          │
          │ 1
          │
          ├──────────────────────┐
          │                      │
          ▼                      ▼
    ┌─────────────┐        ┌─────────────┐
    │PROJET_EQUIPE│        │   RAPPORT   │
    │  (membres)  │        │      *      │
    └─────────────┘        └─────────────┘


Légende:
  ─────── 1   : Cardinalité "un"
  ─────── *   : Cardinalité "plusieurs"
  ─────── 1..* : Un ou plusieurs
  ────────────▶ : Direction de la relation
"""
    para = doc.add_paragraph()
    run = para.add_run(relations_diagram)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    
    doc.add_heading("6.4.2. Description des Relations", 3)
    
    relations_detail = [
        ("Projet → TypeProjet", "ManyToOne", "LAZY", "type_projet_id", "Un projet appartient à un type; un type peut avoir plusieurs projets"),
        ("Projet → Site", "ManyToOne", "LAZY", "site_id", "Un projet est localisé sur un site; un site peut héberger plusieurs projets"),
        ("Projet → Membres", "ElementCollection", "-", "projet_equipe", "Collection de noms de membres de l'équipe projet"),
        ("Rapport → Projet", "ManyToOne", "LAZY", "projet_id", "Un rapport appartient à un projet; un projet peut avoir plusieurs rapports"),
        ("Notification → Utilisateur", "ManyToOne", "LAZY", "utilisateur_id", "Une notification est destinée à un utilisateur"),
        ("Utilisateur → Permissions", "OneToMany", "-", "utilisateur_permissions", "Association des permissions par utilisateur"),
        ("Projet → Phases", "OneToMany", "LAZY", "projet_id", "Un projet est découpé en plusieurs phases séquentielles"),
        ("Phase → Étapes", "OneToMany", "LAZY", "phase_id", "Une phase contient plusieurs étapes à valider"),
        ("Étape → Validations", "OneToMany", "LAZY", "etape_id", "Historique matriciel d'approbation des étapes"),
        ("Projet → Risques", "OneToMany", "LAZY", "projet_id", "Liste des risques identifiés pour le projet"),
        ("Projet → Alertes", "OneToMany", "LAZY", "projet_id", "Alertes temps-réel générées pour surveiller le projet"),
        ("Projet → Historique", "OneToMany", "LAZY", "projet_id", "Piste d'audit de modification des états du projet")
    ]
    create_styled_table(doc, ["Relation", "Type JPA", "Fetch", "Clé Étrangère", "Description"], relations_detail)
    
    doc.add_paragraph()
    
    doc.add_heading("6.4.3. Règles d'Intégrité Référentielle", 3)
    
    integrity_rules = [
        ("fk_rapport_projet", "rapports.projet_id", "projets.id", "CASCADE sur DELETE"),
        ("fk_notification_user", "notifications.utilisateur_id", "utilisateurs.id", "CASCADE sur DELETE"),
        ("fk_projet_type", "projets.type_projet_id", "type_projet.id_type_projet", "SET NULL sur DELETE"),
        ("fk_projet_site", "projets.site_id", "sites.id", "SET NULL sur DELETE")
    ]
    create_styled_table(doc, ["Contrainte", "Table.Colonne", "Référence", "Action"], integrity_rules)
    
    doc.add_page_break()
    
    # 6.5 Dictionnaire de données
    doc.add_heading("6.5. Dictionnaire de Données", 2)
    
    doc.add_heading("6.5.1. Énumérations (ENUM)", 3)
    
    doc.add_paragraph("StatutProjet - États possibles d'un projet:")
    statut_projet = [
        ("PREVU", "À Venir", "Projet planifié mais non démarré"),
        ("EN_COURS", "En Cours", "Projet en cours d'exécution"),
        ("TERMINE", "Terminé", "Projet achevé avec succès"),
        ("ANNULE", "Annulé", "Projet abandonné")
    ]
    create_styled_table(doc, ["Valeur", "Libellé", "Description"], statut_projet)
    
    doc.add_paragraph()
    doc.add_paragraph("PrioriteProjet - Niveaux de priorité:")
    priorite_projet = [
        ("FAIBLE", "Basse", "Priorité basse, traitement différé possible"),
        ("MOYENNE", "Moyenne", "Priorité normale"),
        ("HAUTE", "Haute", "Priorité élevée, attention requise"),
        ("CRITIQUE", "Critique", "Urgence maximale, action immédiate")
    ]
    create_styled_table(doc, ["Valeur", "Libellé", "Description"], priorite_projet)
    
    doc.add_paragraph()
    doc.add_paragraph("NiveauRisque - Évaluation des risques dans les rapports:")
    niveau_risque = [
        ("FAIBLE", "Faible", "Risque minimal, projet sûr"),
        ("MOYEN", "Moyen", "Risques modérés, surveillance requise"),
        ("ELEVE", "Élevé", "Risques significatifs, mesures de mitigation nécessaires"),
        ("CRITIQUE", "Critique", "Risques majeurs, décision stratégique requise")
    ]
    create_styled_table(doc, ["Valeur", "Libellé", "Description"], niveau_risque)
    
    doc.add_paragraph()
    doc.add_paragraph("TypeSite - Catégorisation des sites:")
    type_site = [
        ("SIEGE_SOCIAL", "Siège social principal de l'organisation"),
        ("BUREAU_REGIONAL", "Bureau décentralisé en région"),
        ("CENTRE_OPERATIONNEL", "Centre d'opérations techniques"),
        ("DATACENTER", "Centre de données informatiques"),
        ("SITE_CLIENT", "Emplacement chez un client"),
        ("FORMATION", "Centre de formation")
    ]
    create_styled_table(doc, ["Valeur", "Description"], type_site)
    
    doc.add_paragraph()
    doc.add_paragraph("StatutSite - États possibles d'un site:")
    statut_site = [
        ("ACTIF", "Site opérationnel"),
        ("INACTIF", "Site fermé ou hors service"),
        ("EN_CONSTRUCTION", "Site en cours d'aménagement"),
        ("EN_MAINTENANCE", "Site temporairement indisponible")
    ]
    create_styled_table(doc, ["Valeur", "Description"], statut_site)
    
    doc.add_paragraph()
    doc.add_paragraph("Role - Rôles utilisateur pour l'autorisation:")
    role_enum = [
        ("ADMINISTRATEUR_SYSTEME", "Accès complet, gestion système"),
        ("CHEF_DE_PROJET", "Création et gestion des projets"),
        ("ANALYSTE", "Analyse et consultation des données"),
        ("DEVELOPPEUR", "Accès aux projets techniques"),
        ("DECIDEUR", "Vue stratégique et décisionnelle"),
        ("OBSERVATEUR", "Consultation en lecture seule")
    ]
    create_styled_table(doc, ["Valeur", "Description"], role_enum)
    
    doc.add_page_break()
    
    doc.add_heading("6.5.2. Vues Matérialisées (Statistiques)", 3)
    doc.add_paragraph("L'application utilise des vues pour le calcul des statistiques:")
    
    doc.add_paragraph("vue_statistiques_projets:")
    vue_projets = [
        ("total_projets", "INTEGER", "Nombre total de projets"),
        ("projets_en_cours", "INTEGER", "Projets avec statut EN_COURS"),
        ("projets_termines", "INTEGER", "Projets avec statut TERMINE"),
        ("projets_prevus", "INTEGER", "Projets avec statut PREVU"),
        ("projets_annules", "INTEGER", "Projets avec statut ANNULE"),
        ("budget_total", "DECIMAL", "Somme des budgets alloués"),
        ("progression_moyenne", "DECIMAL", "Moyenne des progressions")
    ]
    create_styled_table(doc, ["Colonne", "Type", "Description"], vue_projets)
    
    doc.add_paragraph()
    doc.add_paragraph("vue_statistiques_rapports:")
    vue_rapports = [
        ("total_rapports", "INTEGER", "Nombre total de rapports"),
        ("rapports_pdf", "INTEGER", "Rapports au format PDF"),
        ("rapports_word", "INTEGER", "Rapports au format Word"),
        ("faisabilite_moyenne", "DECIMAL", "Score moyen de faisabilité"),
        ("risques_faibles", "INTEGER", "Rapports à risque faible"),
        ("risques_moyens", "INTEGER", "Rapports à risque moyen"),
        ("risques_eleves", "INTEGER", "Rapports à risque élevé"),
        ("risques_critiques", "INTEGER", "Rapports à risque critique"),
        ("budget_estime_total", "DECIMAL", "Somme des budgets estimés")
    ]
    create_styled_table(doc, ["Colonne", "Type", "Description"], vue_rapports)
    
    doc.add_page_break()
    
    # === 7. SÉCURITÉ ===
    doc.add_heading("7. Sécurité", 1)
    
    doc.add_heading("7.1. Mesures de Sécurité Implémentées", 2)
    security_measures = [
        ("Authentification JWT", "Tokens signés avec clé secrète 256 bits"),
        ("Mots de passe", "Hachage BCrypt (strength 8)"),
        ("CORS", "Origine restreinte localhost"),
        ("CSRF", "Protection désactivée (API REST stateless)"),
        ("Sessions", "Stateless (pas de stockage côté serveur)"),
        ("Validation", "Bean Validation (JSR-380)"),
        ("SQL Injection", "Protection via JPA/Hibernate"),
        ("XSS", "Sanitization Angular")
    ]
    create_styled_table(doc, ["Mesure", "Détail"], security_measures)
    
    doc.add_heading("7.2. Vulnérabilités Corrigées", 2)
    cves = [
        ("CVE-2025-22235", "Spring Boot", "3.2.8"),
        ("CVE-2024-1597", "PostgreSQL Driver", "42.7.3"),
        ("CVE-2025-48924", "Commons Lang3", "3.14.0")
    ]
    create_styled_table(doc, ["CVE", "Composant", "Version Corrigée"], cves)
    
    doc.add_page_break()
    
    # === 8. RECOMMANDATIONS ===
    doc.add_heading("8. Recommandations", 1)
    
    doc.add_heading("8.1. Améliorations Suggérées", 2)
    improvements = [
        ("Logging", "Implémenter ELK Stack pour centralisation des logs"),
        ("Monitoring", "Ajouter Prometheus/Grafana pour monitoring"),
        ("Cache", "Implémenter Redis pour cache de session et requêtes"),
        ("Tests", "Augmenter la couverture de tests (>80%)"),
        ("CI/CD", "Pipeline Jenkins/GitHub Actions"),
        ("Docker", "Containerisation complète"),
        ("Rate Limiting", "Limiter les requêtes API"),
        ("Audit Log", "Historique complet des actions utilisateur")
    ]
    create_styled_table(doc, ["Domaine", "Recommandation"], improvements)
    
    doc.add_heading("8.2. Évolutions Fonctionnelles Suggérées", 2)
    evolutions = [
        "• Export multi-format (Excel, CSV) - (Fonctionnalité amorcée via ExportExcelService)",
        "• Tableaux de bord personnalisables par l'utilisateur",
        "• Intégration calendrier complet (rappels, délais stricts)",
        "• Module de reporting avancé avec graphiques dynamiques",
        "• Application mobile (Flutter/React Native) pour le personnel terrain",
        "",
        "Note : Les notifications temps réel (WebSockets) et le batch de supervision (Cron) ont déjà été intégrés avec succès au système."
    ]
    for e in evolutions:
        doc.add_paragraph(e)
    
    # Sauvegarde
    output_path = "Etude_Technique_ECarto.docx"
    doc.save(output_path)
    print(f"Document généré avec succès: {output_path}")
    return output_path

if __name__ == "__main__":
    generate_etude_technique()
